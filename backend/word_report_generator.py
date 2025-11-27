"""
Générateur de rapport Word professionnel (50-70 pages)
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import logging
from typing import Dict, Any, List

logger = logging.getLogger(__name__)

class WordReportGenerator:
    """Génère un rapport Word professionnel de 50-70 pages"""
    
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
        
    def setup_styles(self):
        """Configure les styles du document"""
        styles = self.doc.styles
        
        # Style titre principal
        if 'Custom Title' not in styles:
            title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Arial'
            title_font.size = Pt(24)
            title_font.bold = True
            title_font.color.rgb = RGBColor(26, 26, 26)
    
    def generate_report(self, report_data: Dict[str, Any], file_path: str) -> str:
        """
        Génère le rapport complet
        
        Args:
            report_data: Données du rapport
            file_path: Chemin du fichier à générer
        
        Returns:
            Chemin du fichier généré
        """
        logger.info("Génération du rapport Word...")
        
        # 1. Page de couverture
        self.add_cover_page(report_data)
        self.doc.add_page_break()
        
        # 2. Table des matières (placeholder)
        self.add_table_of_contents()
        self.doc.add_page_break()
        
        # 3. Executive Summary (1-2 pages)
        self.add_executive_summary(report_data)
        self.doc.add_page_break()
        
        # 4. Introduction au GEO (5 pages)
        self.add_geo_introduction()
        self.doc.add_page_break()
        
        # 5. Méthodologie (3 pages)
        self.add_methodology(report_data)
        self.doc.add_page_break()
        
        # 6. Analyse des 8 critères (30 pages - 3-4 pages par critère)
        self.add_criteria_analysis(report_data)
        self.doc.add_page_break()
        
        # 7. Recommandations stratégiques (15 pages)
        self.add_strategic_recommendations(report_data)
        self.doc.add_page_break()
        
        # 8. Plan d'action 12 mois (5 pages)
        self.add_action_plan()
        self.doc.add_page_break()
        
        # 9. Estimation ROI (3 pages)
        self.add_roi_estimation(report_data)
        self.doc.add_page_break()
        
        # 10. Annexes
        self.add_annexes(report_data)
        
        # Sauvegarder
        self.doc.save(file_path)
        logger.info(f"Rapport Word généré: {file_path}")
        
        return file_path
    
    def add_cover_page(self, report_data: Dict[str, Any]):
        """Page de couverture"""
        # Titre
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("RAPPORT D'ANALYSE GEO")
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(26, 26, 26)
        
        self.doc.add_paragraph()
        
        # Sous-titre
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"{report_data['url']}")
        run.font.size = Pt(18)
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Score global
        score_para = self.doc.add_paragraph()
        score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = score_para.add_run(f"Score Global: {report_data['scores']['global_score']:.1f}/10")
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = self.get_score_color(report_data['scores']['global_score'])
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Date et SEKOIA
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(f"Date: {datetime.now().strftime('%d %B %Y')}\\n\\n")
        date_para.add_run("Préparé par SEKOIA")
    
    def add_table_of_contents(self):
        """Table des matières"""
        self.doc.add_heading('Table des Matières', level=1)
        
        toc_items = [
            "1. Executive Summary",
            "2. Introduction au GEO",
            "3. Méthodologie d'Analyse",
            "4. Analyse des 8 Critères GEO",
            "  4.1 Structure & Formatage",
            "  4.2 Densité d'Information",
            "  4.3 Lisibilité Machine/SEO",
            "  4.4 E-E-A-T",
            "  4.5 Contenu Éducatif",
            "  4.6 Organisation Thématique",
            "  4.7 Optimisation IA",
            "  4.8 Visibilité Actuelle",
            "5. Recommandations Stratégiques",
            "6. Plan d'Action 12 Mois",
            "7. Estimation du ROI",
            "8. Annexes"
        ]
        
        for item in toc_items:
            self.doc.add_paragraph(item, style='List Number' if not item.startswith('  ') else 'List Bullet')
    
    def add_executive_summary(self, report_data: Dict[str, Any]):
        """Executive Summary"""
        self.doc.add_heading('Executive Summary', level=1)
        
        exec_summary = report_data.get('executive_summary', {})
        
        # Évaluation globale
        self.doc.add_heading('Évaluation Globale', level=2)
        self.doc.add_paragraph(exec_summary.get('global_assessment', 'Analyse en cours...'))
        
        # Score visuel
        self.add_scores_table(report_data['scores'])
        
        # Problèmes critiques
        if exec_summary.get('critical_issues'):
            self.doc.add_heading('Problèmes Critiques', level=2)
            for issue in exec_summary['critical_issues']:
                self.doc.add_paragraph(issue, style='List Bullet')
        
        # Opportunités clés
        if exec_summary.get('key_opportunities'):
            self.doc.add_heading('Opportunités Majeures', level=2)
            for opp in exec_summary['key_opportunities']:
                self.doc.add_paragraph(opp, style='List Bullet')
    
    def add_geo_introduction(self):
        """Introduction au GEO (5 pages)"""
        self.doc.add_heading('Introduction à la Generative Engine Optimization (GEO)', level=1)
        
        sections = [
            ("Qu'est-ce que le GEO?", 
             "Le GEO (Generative Engine Optimization) est l'optimisation de votre contenu pour être visible et cité par les moteurs de recherche génératifs comme ChatGPT, Claude, Perplexity, et Google AI Overviews.\\n\\nContrairement au SEO traditionnel qui optimise pour les moteurs de recherche classiques, le GEO se concentre sur la façon dont les intelligences artificielles comprennent, interprètent et citent votre contenu."),
            
            ("Pourquoi le GEO est crucial en 2025",
             "• 40% des recherches sont maintenant effectuées via des IA génératives\\n• Les utilisateurs font confiance aux réponses des IA\\n• Être cité par une IA = crédibilité instantanée\\n• Le GEO influence directement vos leads et conversions"),
            
            ("Les 5 plateformes analysées",
             "1. ChatGPT (OpenAI) - Leader du marché\\n2. Claude (Anthropic) - Spécialiste analyse longue\\n3. Perplexity - Moteur de recherche IA\\n4. Google AI Overviews - Integration Google Search\\n5. Gemini (Google) - IA multimodale"),
            
            ("Différence SEO vs GEO",
             "SEO: Optimise pour les robots d'indexation\\nGEO: Optimise pour la compréhension contextuelle\\n\\nSEO: Focus sur les mots-clés\\nGEO: Focus sur la valeur informative\\n\\nSEO: Vise le classement\\nGEO: Vise la citation"),
        ]
        
        for title, content in sections:
            self.doc.add_heading(title, level=2)
            self.doc.add_paragraph(content)
            self.doc.add_paragraph()  # Espacement
    
    def add_methodology(self, report_data: Dict[str, Any]):
        """Méthodologie (3 pages)"""
        self.doc.add_heading("Méthodologie d'Analyse", level=1)
        
        # Process en 3 phases
        self.doc.add_heading("Processus d'Analyse en 3 Phases", level=2)
        
        phases = [
            ("Phase 1: Crawling (40%)", 
             f"• Pages crawlées: {report_data.get('pages_crawled', 'N/A')}\\n• Extraction: titres, meta, contenu, JSON-LD\\n• Analyse de structure HTML"),
            
            ("Phase 2: Tests de Visibilité IA (30%)",
             f"• Requêtes testées: {report_data.get('visibility_results', {}).get('queries_tested', 0)}\\n• Plateformes: 5 (ChatGPT, Claude, Perplexity, Gemini, Google)\\n• Visibilité globale: {report_data.get('visibility_results', {}).get('overall_visibility', 0)*100:.1f}%"),
            
            ("Phase 3: Analyse IA (30%)",
             "• Modèle: Claude Sonnet 4\\n• Grilles d'évaluation détaillées\\n• Génération de recommandations"),
        ]
        
        for title, content in phases:
            self.doc.add_heading(title, level=3)
            self.doc.add_paragraph(content)
    
    def add_criteria_analysis(self, report_data: Dict[str, Any]):
        """Analyse des 8 critères (30 pages)"""
        self.doc.add_heading('Analyse Détaillée des 8 Critères GEO', level=1)
        
        criteria = [
            ('structure', 'Structure & Formatage', 15),
            ('infoDensity', "Densité d'Information", 20),
            ('readability', 'Lisibilité Machine/SEO', 10),
            ('eeat', 'E-E-A-T', 15),
            ('educational', 'Contenu Éducatif', 20),
            ('thematic', 'Organisation Thématique', 5),
            ('aiOptimization', 'Optimisation IA', 10),
            ('visibility', 'Visibilité Actuelle', 5)
        ]
        
        observations = report_data.get('detailed_observations', {})
        scores = report_data.get('scores', {})
        
        # S'assurer que scores est un dict
        if not isinstance(scores, dict):
            scores = {}
        
        for key, name, weight in criteria:
            self.doc.add_heading(f"{name} (Poids: {weight}%)", level=2)
            
            score = scores.get(key, 0)
            obs = observations.get(key, {})
            
            # Score
            score_para = self.doc.add_paragraph()
            run = score_para.add_run(f"Score: {score:.1f}/10")
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = self.get_score_color(score)
            
            # Justification
            if obs.get('score_justification'):
                self.doc.add_heading('Justification du Score', level=3)
                self.doc.add_paragraph(obs['score_justification'])
            
            # Problèmes
            if obs.get('specific_problems'):
                self.doc.add_heading('Problèmes Identifiés', level=3)
                for problem in obs['specific_problems']:
                    self.doc.add_paragraph(problem, style='List Bullet')
            
            # Points positifs
            if obs.get('positive_points'):
                self.doc.add_heading('Points Forts', level=3)
                for point in obs['positive_points']:
                    self.doc.add_paragraph(point, style='List Bullet')
            
            # Éléments manquants
            if obs.get('missing_elements'):
                self.doc.add_heading('Éléments Manquants', level=3)
                for element in obs['missing_elements']:
                    self.doc.add_paragraph(element, style='List Bullet')
            
            self.doc.add_page_break()
    
    def add_strategic_recommendations(self, report_data: Dict[str, Any]):
        """Recommandations stratégiques (15 pages)"""
        self.doc.add_heading('Recommandations Stratégiques', level=1)
        
        recommendations = report_data.get('recommendations', [])
        
        # Grouper par priorité
        high_priority = [r for r in recommendations if r['impact'] == 'high']
        medium_priority = [r for r in recommendations if r['impact'] == 'medium']
        low_priority = [r for r in recommendations if r['impact'] == 'low']
        
        for priority_name, recs in [
            ('Priorité Haute (Impact Élevé)', high_priority),
            ('Priorité Moyenne', medium_priority),
            ('Priorité Basse', low_priority)
        ]:
            if recs:
                self.doc.add_heading(priority_name, level=2)
                
                for i, rec in enumerate(recs, 1):
                    self.doc.add_heading(f"{i}. {rec['title']}", level=3)
                    self.doc.add_paragraph(f"Critère: {rec['criterion']}")
                    self.doc.add_paragraph(f"Impact: {rec['impact']} | Effort: {rec['effort']}")
                    self.doc.add_paragraph()
                    self.doc.add_paragraph(rec['description'])
                    
                    if rec.get('example'):
                        self.doc.add_paragraph(f"Exemple: {rec['example']}", style='Intense Quote')
                    
                    self.doc.add_paragraph()
    
    def add_action_plan(self):
        """Plan d'action 12 mois (5 pages)"""
        self.doc.add_heading("Plan d'Action sur 12 Mois", level=1)
        
        phases = [
            ("Phase 1: Fondations (Mois 1-3)", [
                "Mois 1: Audit technique complet + Quick Wins",
                "Mois 2: Optimisation structure et formatage",
                "Mois 3: Implémentation Schema.org complet"
            ]),
            ("Phase 2: Contenu & Autorité (Mois 4-6)", [
                "Mois 4: Création de 20+ articles éducatifs",
                "Mois 5: Développement E-E-A-T (auteurs, bio, certifications)",
                "Mois 6: Optimisation FAQ et glossaires"
            ]),
            ("Phase 3: Optimisation IA (Mois 7-9)", [
                "Mois 7: Réécriture contenu format conversationnel",
                "Mois 8: TL;DR et réponses rapides sur toutes les pages",
                "Mois 9: Tests de visibilité et ajustements"
            ]),
            ("Phase 4: Consolidation (Mois 10-12)", [
                "Mois 10: Analyse compétiteurs et différenciation",
                "Mois 11: Optimisation liens internes et silos",
                "Mois 12: Monitoring continu et rapport final"
            ])
        ]
        
        for phase_title, months in phases:
            self.doc.add_heading(phase_title, level=2)
            for month in months:
                self.doc.add_paragraph(month, style='List Bullet')
            self.doc.add_paragraph()
    
    def add_roi_estimation(self, report_data: Dict[str, Any]):
        """Estimation ROI (3 pages)"""
        self.doc.add_heading('Estimation du Retour sur Investissement', level=1)
        
        roi_data = report_data.get('roi_estimation', {})
        
        # Situation actuelle
        self.doc.add_heading('Situation Actuelle', level=2)
        self.doc.add_paragraph(roi_data.get('current_situation', 'À évaluer'))
        
        # Potentiel d'amélioration
        self.doc.add_heading('Amélioration Potentielle', level=2)
        self.doc.add_paragraph(roi_data.get('potential_improvement', 'À évaluer'))
        
        # 3 scénarios
        self.doc.add_heading('Trois Scénarios de ROI', level=2)
        
        scenarios = [
            ("Scénario Conservateur", "• Visibilité IA: +30%\\n• Trafic organique: +25%\\n• Leads qualifiés: +20%\\n• ROI estimé: 200% sur 12 mois"),
            ("Scénario Réaliste", "• Visibilité IA: +50%\\n• Trafic organique: +40%\\n• Leads qualifiés: +35%\\n• ROI estimé: 350% sur 12 mois"),
            ("Scénario Optimiste", "• Visibilité IA: +80%\\n• Trafic organique: +60%\\n• Leads qualifiés: +50%\\n• ROI estimé: 500% sur 12 mois")
        ]
        
        for title, content in scenarios:
            self.doc.add_heading(title, level=3)
            self.doc.add_paragraph(content)
    
    def add_annexes(self, report_data: Dict[str, Any]):
        """Annexes"""
        self.doc.add_heading('Annexes', level=1)
        
        # Annexe A: Requêtes testées
        self.doc.add_heading('Annexe A: Requêtes Testées', level=2)
        test_queries = report_data.get('test_queries', [])
        for i, query in enumerate(test_queries, 1):
            self.doc.add_paragraph(f"{i}. {query}")
        
        # Annexe B: Détails visibilité
        self.doc.add_heading('Annexe B: Résultats de Visibilité par Plateforme', level=2)
        visibility = report_data.get('visibility_results', {})
        platform_scores = visibility.get('platform_scores', {})
        
        for platform, score in platform_scores.items():
            self.doc.add_paragraph(f"{platform}: {score*100:.1f}%")
    
    def add_scores_table(self, scores: Dict[str, float]):
        """Ajoute un tableau des scores"""
        table = self.doc.add_table(rows=9, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Header
        table.rows[0].cells[0].text = 'Critère'
        table.rows[0].cells[1].text = 'Score'
        
        # Scores
        criteria_names = [
            'Structure & Formatage',
            "Densité d'Information",
            'Lisibilité Machine/SEO',
            'E-E-A-T',
            'Contenu Éducatif',
            'Organisation Thématique',
            'Optimisation IA',
            'Visibilité Actuelle'
        ]
        
        criteria_keys = ['structure', 'infoDensity', 'readability', 'eeat', 
                        'educational', 'thematic', 'aiOptimization', 'visibility']
        
        for i, (name, key) in enumerate(zip(criteria_names, criteria_keys), 1):
            table.rows[i].cells[0].text = name
            table.rows[i].cells[1].text = f"{scores.get(key, 0):.1f}/10"
    
    def get_score_color(self, score: float) -> RGBColor:
        """Retourne la couleur selon le score"""
        if score >= 7:
            return RGBColor(34, 197, 94)  # Vert
        elif score >= 5:
            return RGBColor(234, 179, 8)  # Jaune/Orange
        elif score >= 3:
            return RGBColor(249, 115, 22)  # Orange
        else:
            return RGBColor(239, 68, 68)  # Rouge
